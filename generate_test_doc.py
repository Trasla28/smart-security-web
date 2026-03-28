"""Genera el documento Word de pruebas del aplicativo Smart Security Tickets."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Paleta de colores ─────────────────────────────────────────────────────────
AZUL_OSCURO   = RGBColor(0x1A, 0x2C, 0x4E)   # #1a2c4e
AZUL_MEDIO    = RGBColor(0x1E, 0x40, 0x7A)   # encabezados
AZUL_CLARO    = RGBColor(0xDB, 0xEA, 0xFE)   # fondo celdas encabezado
VERDE         = RGBColor(0x16, 0x65, 0x34)   # resultado esperado
ROJO          = RGBColor(0x99, 0x1B, 0x1B)   # advertencias
GRIS_CLARO    = RGBColor(0xF3, 0xF4, 0xF6)   # filas alternas
NEGRO         = RGBColor(0x11, 0x18, 0x27)
BLANCO        = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Configuración de página ──────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.59)
section.page_height = Cm(27.94)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="D1D5DB", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def heading1(text: str):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = BLANCO
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.bold = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1A2C4E")
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p


def heading2(text: str):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = AZUL_OSCURO
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def heading3(text: str):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = AZUL_MEDIO
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p


def body(text: str, bold=False, italic=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(text: str, bold_prefix: str = None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def info_box(text: str, bg="EFF6FF", border="BFDBFE", text_color=None):
    """Caja informativa con fondo de color."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_borders(cell, border, 6)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    if text_color:
        run.font.color.rgb = text_color
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def evidence_box():
    """Caja de evidencia vacía para pegar screenshot."""
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_bg(cell, "F9FAFB")
    set_cell_borders(cell, "9CA3AF", 8)
    p = cell.paragraphs[0]
    run = p.add_run("[ Pegar evidencia / captura de pantalla aquí ]")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.height = Cm(6)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def result_row(label: str, expected: str):
    """Fila compacta: Resultado esperado."""
    p = doc.add_paragraph()
    r1 = p.add_run("✔ Resultado esperado: ")
    r1.font.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = VERDE
    r2 = p.add_run(expected)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
    p.paragraph_format.space_after = Pt(2)


def step_table(steps: list[tuple]):
    """
    Tabla de pasos: (N°, Acción, Resultado esperado)
    steps = [(num, accion, resultado), ...]
    """
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Anchos
    t.columns[0].width = Cm(1.2)
    t.columns[1].width = Cm(8.5)
    t.columns[2].width = Cm(7.3)

    # Encabezado
    hdr = t.rows[0].cells
    for cell, txt, bg in zip(hdr, ["Paso", "Acción", "Resultado esperado"], ["1A2C4E","1A2C4E","1A2C4E"]):
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BLANCO
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Filas
    for i, (num, accion, resultado) in enumerate(steps):
        row = t.add_row()
        bg = "FFFFFF" if i % 2 == 0 else "F8FAFC"

        cell_num = row.cells[0]
        set_cell_bg(cell_num, bg)
        p = cell_num.paragraphs[0]
        p.add_run(str(num)).font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_act = row.cells[1]
        set_cell_bg(cell_act, bg)
        p = cell_act.paragraphs[0]
        p.add_run(accion).font.size = Pt(9)

        cell_res = row.cells[2]
        set_cell_bg(cell_res, bg)
        p = cell_res.paragraphs[0]
        run = p.add_run(resultado)
        run.font.size = Pt(9)
        run.font.color.rgb = VERDE

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def cred_table(rows_data: list[tuple]):
    """Tabla de credenciales."""
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for cell, txt in zip(t.rows[0].cells, ["Rol", "Email", "Contraseña", "Permisos"]):
        set_cell_bg(cell, "1A2C4E")
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BLANCO
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bgs = ["FFFFFF", "F0F9FF", "FFFFFF", "F0F9FF", "FFFFFF"]
    for i, (rol, email, pwd, perms) in enumerate(rows_data):
        row = t.add_row()
        bg = bgs[i % len(bgs)]
        for cell, txt in zip(row.cells, [rol, email, pwd, perms]):
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.add_run(txt).font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return t


# ════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Cm(3)
p.paragraph_format.space_after  = Pt(0)

# Logo texto SS
run = p.add_run("SS")
run.font.size = Pt(40)
run.font.bold = True
run.font.color.rgb = AZUL_OSCURO
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc.add_paragraph()
run2 = p2.add_run("Smart Security Tickets")
run2.font.size = Pt(28)
run2.font.bold = True
run2.font.color.rgb = AZUL_OSCURO
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

p3 = doc.add_paragraph()
run3 = p3.add_run("GUÍA DE PRUEBAS DEL APLICATIVO")
run3.font.size = Pt(14)
run3.font.bold = True
run3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph().paragraph_format.space_after = Pt(20)

# Cuadro de datos del documento
t_meta = doc.add_table(rows=5, cols=2)
t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
t_meta.style = "Table Grid"
meta = [
    ("Versión", "1.0"),
    ("Fecha", datetime.date.today().strftime("%d de %B de %Y")),
    ("Preparado por", "Equipo de Desarrollo"),
    ("URL de pruebas", "http://localhost:3000"),
    ("Estado", "Para revisión y aprobación"),
]
for i, (k, v) in enumerate(meta):
    bg = "EFF6FF" if i % 2 == 0 else "FFFFFF"
    c0, c1 = t_meta.rows[i].cells
    set_cell_bg(c0, "1E407A"); set_cell_bg(c1, bg)
    r0 = c0.paragraphs[0].add_run(k); r0.font.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = BLANCO
    r1 = c1.paragraphs[0].add_run(v); r1.font.size = Pt(9)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 1 — CREDENCIALES
# ════════════════════════════════════════════════════════════════════════════

heading1("1. CREDENCIALES DE ACCESO")
body("Todos los usuarios usan la misma URL de ingreso:", size=10)
info_box("🌐  URL del sistema: http://localhost:3000\n   El campo Tenant se rellena automáticamente, no es necesario modificarlo.", "EFF6FF", "93C5FD")

body("La siguiente tabla contiene todos los usuarios disponibles para las pruebas:", size=10)

cred_table([
    ("Administrador",  "admin@smartsecurity.com.co",      "Admin123!",  "Acceso total al sistema"),
    ("Supervisor",     "supervisor@smartsecurity.com.co",  "Super123!",  "Reportes, gestión de equipo"),
    ("Agente TI",      "agente.ti@smartsecurity.com.co",   "Agent123!",  "Atender y cerrar tickets asignados"),
    ("Agente RRHH",    "agente.rrhh@smartsecurity.com.co", "Agent123!",  "Atender y cerrar tickets asignados"),
    ("Solicitante",    "solicitante@smartsecurity.com.co",  "User123!",   "Crear y seguir sus tickets"),
])

info_box("⚠  En cada bloque de pruebas se indica con qué usuario debes estar conectado antes de ejecutar los pasos.", "FEF9C3", "FCD34D", ROJO)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 2 — AUTENTICACIÓN
# ════════════════════════════════════════════════════════════════════════════

heading1("2. PRUEBAS DE AUTENTICACIÓN")

heading2("2.1 Login exitoso")
body("Usuario: cualquiera de la tabla anterior", bold=True)

step_table([
    (1, "Abrir http://localhost:3000 en el navegador",
       "Se muestra la pantalla de login con logo SS, campo email y contraseña"),
    (2, "Ingresar email: admin@smartsecurity.com.co  /  Contraseña: Admin123! y hacer clic en 'Ingresar'",
       "El sistema redirige al Dashboard principal"),
    (3, "Verificar que en la barra lateral aparece el nombre del usuario y el rol 'Administrador'",
       "Nombre e icono de rol visibles en el sidebar"),
    (4, "Hacer clic en el botón de cerrar sesión (ícono en la barra superior derecha)",
       "El sistema redirige de nuevo a la pantalla de login"),
])
evidence_box()

heading2("2.2 Login fallido")
step_table([
    (1, "En la pantalla de login, ingresar email: admin@smartsecurity.com.co  /  Contraseña: incorrecta123",
       "Aparece mensaje de error: 'Credenciales inválidas. Verifica tu email y contraseña.'"),
    (2, "Intentar acceder directamente a http://localhost:3000/tickets sin estar autenticado",
       "El sistema redirige automáticamente a /login"),
])
evidence_box()

heading2("2.3 Verificar acceso por rol (Solicitante no ve Admin)")
body("Iniciar sesión con: solicitante@smartsecurity.com.co / User123!", bold=True)
step_table([
    (1, "Iniciar sesión con el usuario Solicitante",
       "Dashboard visible, sidebar muestra: Dashboard, Mis Tickets"),
    (2, "Intentar navegar a http://localhost:3000/admin/users",
       "El sistema redirige a /login o muestra pantalla de acceso denegado. El menú 'Admin' NO aparece en el sidebar"),
    (3, "Cerrar sesión",
       "Regresa a pantalla de login"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 3 — DASHBOARD
# ════════════════════════════════════════════════════════════════════════════

heading1("3. PRUEBAS DE DASHBOARD")
body("Usuario: admin@smartsecurity.com.co / Admin123!", bold=True)

heading2("3.1 Visualización del Dashboard")
step_table([
    (1, "Iniciar sesión como Administrador y navegar a la página principal (ícono 'Dashboard' en sidebar)",
       "Se muestran 4 tarjetas: Abiertos, En Proceso, Resueltos hoy, % SLA"),
    (2, "Verificar que se muestra la gráfica 'Tickets por área' (barras horizontales)",
       "Gráfica visible con datos de las áreas: Tecnología, Recursos Humanos, Operaciones"),
    (3, "Verificar la gráfica 'Distribución por estado' (donut)",
       "Gráfica de dona con colores por estado visible"),
    (4, "Verificar la tabla 'Desempeño por agente' en la parte inferior",
       "Tabla con columnas: Agente, Asignados, Resueltos, Prom. resolución, Cumpl. SLA"),
])
evidence_box()

heading2("3.2 Dashboard como Supervisor")
body("Cerrar sesión. Iniciar sesión con: supervisor@smartsecurity.com.co / Super123!", bold=True)
step_table([
    (1, "Iniciar sesión como Supervisor y abrir el Dashboard",
       "Se muestran las mismas estadísticas y gráficas que el admin"),
    (2, "Verificar que la tabla de desempeño por agente también es visible",
       "Tabla visible (supervisor también tiene acceso)"),
])
evidence_box()

heading2("3.3 Dashboard como Agente (vista limitada)")
body("Cerrar sesión. Iniciar sesión con: agente.ti@smartsecurity.com.co / Agent123!", bold=True)
step_table([
    (1, "Iniciar sesión como Agente TI y abrir el Dashboard",
       "Se muestran las 4 tarjetas de estadísticas y las 2 gráficas"),
    (2, "Verificar que la tabla de desempeño por agente NO aparece",
       "La tabla de agentes no es visible para el rol Agente"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 4 — GESTIÓN DE TICKETS (SOLICITANTE)
# ════════════════════════════════════════════════════════════════════════════

heading1("4. PRUEBAS DE TICKETS — SOLICITANTE")
body("Usuario: solicitante@smartsecurity.com.co / User123!", bold=True)

heading2("4.1 Crear un ticket nuevo")
step_table([
    (1,  "Iniciar sesión como Solicitante",
         "Dashboard visible"),
    (2,  "Hacer clic en 'Tickets' en el sidebar y luego en el botón '+ Nuevo Ticket'",
         "Se abre el formulario de creación de ticket"),
    (3,  "Ingresar en Título: 'Mi computador no enciende'",
         "Campo lleno sin errores"),
    (4,  "Ingresar en Descripción: 'Al presionar el botón de encendido no hay ninguna respuesta, ni luces ni sonidos.'",
         "Campo lleno sin errores"),
    (5,  "Seleccionar Prioridad: Alta",
         "Prioridad Alta seleccionada"),
    (6,  "Seleccionar Categoría: Hardware (o la disponible)",
         "Categoría seleccionada"),
    (7,  "Hacer clic en 'Crear Ticket'",
         "Aparece mensaje de éxito y se redirige al detalle del ticket con número #TK-XXXX asignado"),
])
evidence_box()

heading2("4.2 Ver la lista de tickets")
step_table([
    (1, "Hacer clic en 'Tickets' en el sidebar",
        "Se muestra la tabla de tickets con el ticket recién creado"),
    (2, "Verificar las columnas visibles: N°, Título, Estado, Prioridad, Área, Agente asignado, SLA, Fecha",
        "Todas las columnas visibles"),
    (3, "Usar el filtro de Estado: seleccionar 'Abierto'",
        "La tabla muestra solo los tickets en estado 'Abierto'"),
    (4, "Limpiar el filtro",
        "La tabla vuelve a mostrar todos los tickets"),
])
evidence_box()

heading2("4.3 Ver el detalle de un ticket")
step_table([
    (1, "Hacer clic en el ticket 'Mi computador no enciende'",
        "Se abre la vista de detalle del ticket"),
    (2, "Verificar que se muestra: título, descripción, estado, prioridad, solicitante, fechas",
        "Todos los campos visibles en la columna izquierda"),
    (3, "Hacer clic en la pestaña 'Comentarios'",
        "Se muestra el área de comentarios (vacía inicialmente)"),
    (4, "Escribir un comentario: 'Esto es urgente, necesito el equipo para trabajar' y hacer clic en Enviar",
        "El comentario aparece en la lista con la hora y nombre del usuario"),
    (5, "Hacer clic en la pestaña 'Historial'",
        "Se muestra el historial con el evento 'Ticket creado'"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 5 — GESTIÓN DE TICKETS (AGENTE)
# ════════════════════════════════════════════════════════════════════════════

heading1("5. PRUEBAS DE TICKETS — AGENTE TI")
body("Cerrar sesión. Iniciar sesión con: agente.ti@smartsecurity.com.co / Agent123!", bold=True)

heading2("5.1 Ver tickets asignados y tomar uno")
step_table([
    (1,  "Iniciar sesión como Agente TI y navegar a Tickets",
         "Se muestra la lista de tickets (los tickets del tenant son visibles para el agente)"),
    (2,  "Hacer clic en el ticket 'Mi computador no enciende'",
         "Se abre el detalle del ticket en estado 'Abierto'"),
    (3,  "Hacer clic en el botón 'En Proceso' (si está disponible) o verificar los botones de acción",
         "El estado del ticket cambia a 'En Proceso'"),
    (4,  "Verificar que en el historial aparece el cambio de estado",
         "Historial registra: 'Estado cambiado de Abierto a En Proceso'"),
])
evidence_box()

heading2("5.2 Agregar comentario interno")
step_table([
    (1,  "Dentro del ticket, escribir en el campo de comentario: 'Estoy revisando el hardware del equipo'",
         "Campo de texto disponible"),
    (2,  "Activar la opción 'Nota interna' (si existe el toggle) antes de enviar",
         "El checkbox/toggle de nota interna está marcado"),
    (3,  "Hacer clic en Enviar",
         "El comentario aparece con estilo diferente (fondo amarillo/anaranjado) indicando que es interno"),
])
evidence_box()

heading2("5.3 Resolver un ticket")
step_table([
    (1,  "Dentro del mismo ticket, hacer clic en el botón 'Resolver'",
         "El estado del ticket cambia a 'Resuelto'"),
    (2,  "Verificar el historial del ticket",
         "El historial muestra el cambio a 'Resuelto' con fecha y hora"),
    (3,  "Navegar a la lista de tickets y verificar que el estado aparece actualizado",
         "El ticket muestra badge verde 'Resuelto'"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 6 — GESTIÓN DE TICKETS (SOLICITANTE — REABRIR)
# ════════════════════════════════════════════════════════════════════════════

heading1("6. PRUEBAS DE TICKETS — REABRIR Y CERRAR")

heading2("6.1 Solicitante reabre un ticket resuelto")
body("Cerrar sesión. Iniciar sesión con: solicitante@smartsecurity.com.co / User123!", bold=True)
step_table([
    (1,  "Iniciar sesión como Solicitante y abrir el ticket 'Mi computador no enciende'",
         "El ticket muestra estado 'Resuelto' y aparece el botón 'Reabrir'"),
    (2,  "Hacer clic en 'Reabrir'",
         "Aparece campo para ingresar el motivo"),
    (3,  "Escribir: 'El problema persiste, el equipo sigue sin encender' y confirmar",
         "El ticket vuelve al estado 'Abierto' o 'En Proceso', el historial registra el reapertura"),
])
evidence_box()

heading2("6.2 Agente cierra manualmente un ticket")
body("Cerrar sesión. Iniciar sesión con: agente.ti@smartsecurity.com.co / Agent123!", bold=True)
step_table([
    (1,  "Abrir el ticket y cambiar estado a 'Resuelto' nuevamente",
         "Estado actualizado a Resuelto"),
    (2,  "Hacer clic en el botón 'Cerrar'",
         "El ticket pasa a estado 'Cerrado'"),
    (3,  "Verificar que en estado 'Cerrado' ya no aparecen botones de cambio de estado para el agente",
         "Los botones de acción desaparecen o están deshabilitados"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 7 — ADMINISTRACIÓN DE USUARIOS
# ════════════════════════════════════════════════════════════════════════════

heading1("7. PRUEBAS DE ADMINISTRACIÓN — USUARIOS")
body("Cerrar sesión. Iniciar sesión con: admin@smartsecurity.com.co / Admin123!", bold=True)

heading2("7.1 Crear un usuario nuevo")
step_table([
    (1,  "Ir a Admin → Usuarios en el sidebar",
         "Se muestra la lista de los 5 usuarios existentes"),
    (2,  "Hacer clic en '+ Nuevo Usuario'",
         "Se abre el modal de creación de usuario"),
    (3,  "Ingresar: Nombre: Pedro Prueba  |  Email: pedro@smartsecurity.com.co  |  Rol: Agente  |  Contraseña: Test123!",
         "Campos llenos sin errores"),
    (4,  "Hacer clic en 'Guardar' o 'Crear'",
         "El modal se cierra y el usuario 'Pedro Prueba' aparece en la lista"),
])
evidence_box()

heading2("7.2 Editar un usuario existente")
step_table([
    (1,  "Hacer clic en el botón de editar del usuario 'Pedro Prueba'",
         "Se abre el modal con los datos del usuario precargados"),
    (2,  "Cambiar el Rol a 'Supervisor' y guardar",
         "El modal se cierra y la tabla muestra el rol actualizado para Pedro Prueba"),
])
evidence_box()

heading2("7.3 Archivar un usuario")
step_table([
    (1,  "Hacer clic en el botón de archivar del usuario 'Pedro Prueba'",
         "Aparece confirmación: '¿Estás seguro de archivar este usuario?'"),
    (2,  "Confirmar la acción",
         "El usuario desaparece de la lista activa (o aparece con estado 'Archivado')"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 8 — ADMINISTRACIÓN DE ÁREAS
# ════════════════════════════════════════════════════════════════════════════

heading1("8. PRUEBAS DE ADMINISTRACIÓN — ÁREAS")
body("Usuario: admin@smartsecurity.com.co / Admin123! (continuar con la sesión activa)", bold=True)

heading2("8.1 Ver áreas existentes")
step_table([
    (1,  "Ir a Admin → Áreas en el sidebar",
         "Se muestran las áreas: Tecnología, Recursos Humanos, Operaciones"),
    (2,  "Expandir el panel de un área haciendo clic en ella",
         "Se muestran los miembros del área"),
])
evidence_box()

heading2("8.2 Crear un área nueva")
step_table([
    (1,  "Hacer clic en '+ Nueva Área'",
         "Se abre formulario/modal de creación"),
    (2,  "Ingresar Nombre: 'Mantenimiento'  y  Descripción: 'Área de mantenimiento de instalaciones'",
         "Campos llenos"),
    (3,  "Guardar",
         "El área 'Mantenimiento' aparece en la lista"),
])
evidence_box()

heading2("8.3 Agregar miembro a un área")
step_table([
    (1,  "Expandir el área 'Tecnología'",
         "Se muestran los miembros actuales"),
    (2,  "Usar la opción de agregar miembro y seleccionar 'Agente RRHH'",
         "El agente RRHH aparece como miembro del área Tecnología"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 9 — ADMINISTRACIÓN DE CATEGORÍAS Y SLAs
# ════════════════════════════════════════════════════════════════════════════

heading1("9. PRUEBAS DE ADMINISTRACIÓN — CATEGORÍAS Y SLAs")
body("Usuario: admin@smartsecurity.com.co / Admin123!", bold=True)

heading2("9.1 Ver y editar categorías")
step_table([
    (1,  "Ir a Admin → Categorías",
         "Se muestran las categorías: Hardware, Software, Nómina, Vacaciones, General"),
    (2,  "Hacer clic en editar sobre la categoría 'General'",
         "Se abre el modal con los datos de la categoría"),
    (3,  "Cambiar el área por defecto a 'Tecnología' y guardar",
         "La categoría 'General' ahora muestra área 'Tecnología'"),
    (4,  "Hacer clic en el toggle de activar/desactivar de la categoría 'General'",
         "El estado de la categoría cambia (activa ↔ inactiva)"),
])
evidence_box()

heading2("9.2 Ver SLAs")
step_table([
    (1,  "Ir a Admin → SLAs",
         "Se muestran los 4 SLAs: Urgente (1h/4h), Alta (4h/8h), Media (8h/24h), Baja (24h/72h)"),
    (2,  "Hacer clic en editar sobre el SLA 'Media'",
         "Se abre el modal con las horas configuradas"),
    (3,  "Cambiar el tiempo de respuesta a 6 horas y guardar",
         "El SLA Media ahora muestra 6h de respuesta en la tabla"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 10 — PLANTILLAS RECURRENTES
# ════════════════════════════════════════════════════════════════════════════

heading1("10. PRUEBAS DE PLANTILLAS RECURRENTES")
body("Usuario: admin@smartsecurity.com.co / Admin123!", bold=True)

heading2("10.1 Crear una plantilla recurrente")
step_table([
    (1,  "Ir a Admin → Plantillas Recurrentes",
         "Se muestra la lista (vacía o con plantillas existentes)"),
    (2,  "Hacer clic en '+ Nueva Plantilla'",
         "Se abre el modal de creación"),
    (3,  "Configurar: Título: 'Respaldo semanal de datos'  |  Descripción: 'Verificar que los respaldos automáticos se ejecutaron correctamente'  |  Prioridad: Media  |  Recurrencia: Semanal  |  Día: Lunes",
         "Campos configurados"),
    (4,  "Guardar la plantilla",
         "La plantilla aparece en la lista con estado 'Activa' y la fecha del próximo ticket"),
    (5,  "Usar el toggle para desactivar la plantilla",
         "La plantilla muestra estado 'Inactiva'"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 11 — CONFIGURACIÓN DEL TENANT
# ════════════════════════════════════════════════════════════════════════════

heading1("11. PRUEBAS DE CONFIGURACIÓN DEL TENANT")
body("Usuario: admin@smartsecurity.com.co / Admin123!", bold=True)

heading2("11.1 Ver y editar la configuración")
step_table([
    (1,  "Ir a Admin → Configuración",
         "Se muestra el formulario de configuración con los valores actuales del tenant"),
    (2,  "Verificar que se muestran los campos: Color primario, Método de autenticación, Días de auto-cierre, Horario laboral, Días laborales, Reporte semanal",
         "Todos los campos visibles con sus valores actuales"),
    (3,  "Cambiar el color primario a '#2563EB' usando el selector de color",
         "El selector de color muestra el nuevo color"),
    (4,  "Cambiar el umbral de abuso de urgencia a 30",
         "Campo actualizado a 30"),
    (5,  "Hacer clic en 'Guardar configuración'",
         "Aparece mensaje de éxito: 'Configuración guardada'"),
    (6,  "Recargar la página (F5) y verificar que los valores persisten",
         "Los valores guardados se mantienen tras recargar"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 12 — REPORTES Y MÉTRICAS
# ════════════════════════════════════════════════════════════════════════════

heading1("12. PRUEBAS DE REPORTES")
body("Cerrar sesión. Iniciar sesión con: supervisor@smartsecurity.com.co / Super123!", bold=True)

heading2("12.1 Ver reporte de abuso de urgencia")
step_table([
    (1,  "Iniciar sesión como Supervisor y hacer clic en 'Reportes' en el sidebar",
         "Se muestra la página de reportes"),
    (2,  "Verificar que se muestra la tabla de 'Abuso de urgencia'",
         "Tabla visible con columnas: Solicitante, Tickets urgentes, % del total"),
    (3,  "Verificar que los datos son coherentes con los tickets de prueba del seed",
         "Se muestran datos de los solicitantes con tickets urgentes"),
])
evidence_box()

heading2("12.2 Acceso denegado al reporte como Agente")
body("Cerrar sesión. Iniciar sesión con: agente.ti@smartsecurity.com.co / Agent123!", bold=True)
step_table([
    (1,  "Iniciar sesión como Agente TI",
         "El sidebar NO muestra la opción 'Reportes'"),
    (2,  "Intentar navegar directamente a http://localhost:3000/reports",
         "El sistema redirige a login o muestra acceso denegado"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 13 — NOTIFICACIONES
# ════════════════════════════════════════════════════════════════════════════

heading1("13. PRUEBAS DE NOTIFICACIONES")

heading2("13.1 Recibir notificación al asignar un ticket")
body("Para esta prueba necesitas dos navegadores o una ventana privada:", bold=True)
body("  • Navegador A: agente.ti@smartsecurity.com.co / Agent123!")
body("  • Navegador B: admin@smartsecurity.com.co / Admin123!")

step_table([
    (1,  "[Navegador B - Admin] Ir a la lista de tickets y abrir cualquier ticket sin asignar",
         "Detalle del ticket visible"),
    (2,  "[Navegador B - Admin] Asignar el ticket al 'Agente TI' usando el botón de asignación",
         "El ticket muestra 'Agente TI' como asignado"),
    (3,  "[Navegador A - Agente TI] Observar el ícono de campana en la barra superior",
         "Aparece un punto rojo/número en la campana indicando notificación nueva"),
    (4,  "[Navegador A - Agente TI] Hacer clic en la campana",
         "Se despliega lista de notificaciones con: 'Te asignaron el ticket #TK-XXXX'"),
    (5,  "[Navegador A - Agente TI] Hacer clic en 'Marcar todo como leído'",
         "El punto rojo desaparece de la campana"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 14 — FLUJO COMPLETO E2E
# ════════════════════════════════════════════════════════════════════════════

heading1("14. FLUJO COMPLETO END-TO-END")
body("Esta prueba simula el ciclo de vida completo de un ticket desde la creación hasta el cierre.", size=10)
info_box("Ejecutar esta prueba al final, luego de haber probado las secciones anteriores. Usar ventanas privadas para cambiar de usuario fácilmente.", "F0FDF4", "86EFAC")

step_table([
    (1,  "[Solicitante] Crear ticket: 'Internet sin conexión en sala de reuniones 2'  |  Prioridad: Alta  |  Categoría: Hardware",
         "Ticket creado con número #TK-XXXX y estado 'Abierto'"),
    (2,  "[Agente TI] Abrir el ticket y cambiar estado a 'En Proceso'",
         "Estado 'En Proceso', historial actualizado"),
    (3,  "[Agente TI] Agregar comentario público: 'Estoy revisando la configuración del switch de red'",
         "Comentario visible para el solicitante"),
    (4,  "[Solicitante] Verificar que el comentario del agente es visible en el ticket",
         "Comentario visible con el nombre del agente"),
    (5,  "[Agente TI] Cambiar estado a 'Resuelto'",
         "Estado 'Resuelto', notificación enviada al solicitante"),
    (6,  "[Solicitante] Revisar el ticket y hacer clic en 'Reabrir' con motivo: 'La conexión falló de nuevo'",
         "Ticket vuelve a 'Abierto', historial registra la reapertura"),
    (7,  "[Agente TI] Volver a resolver el ticket y luego cerrarlo",
         "Ticket en estado 'Cerrado'"),
    (8,  "[Admin] Verificar en el Dashboard que las estadísticas reflejan el ticket cerrado",
         "Las tarjetas y gráficas del dashboard muestran datos actualizados"),
])
evidence_box()

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 15 — MATRIZ DE RESULTADOS
# ════════════════════════════════════════════════════════════════════════════

heading1("15. MATRIZ DE RESULTADOS")
body("Completar esta tabla al finalizar todas las pruebas:", size=10)

# Tabla de resultados
t_res = doc.add_table(rows=1, cols=5)
t_res.style = "Table Grid"
t_res.alignment = WD_TABLE_ALIGNMENT.LEFT

for cell, txt in zip(t_res.rows[0].cells, ["Sección", "Descripción", "Resultado\n(✔/✗/N/A)", "Observaciones", "Evidencia"]):
    set_cell_bg(cell, "1A2C4E")
    p = cell.paragraphs[0]
    run = p.add_run(txt); run.font.bold = True; run.font.size = Pt(8.5); run.font.color.rgb = BLANCO
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

secciones = [
    ("2.1", "Login exitoso"),
    ("2.2", "Login fallido"),
    ("2.3", "Control de acceso por rol"),
    ("3.1", "Dashboard — Admin"),
    ("3.2", "Dashboard — Supervisor"),
    ("3.3", "Dashboard — Agente"),
    ("4.1", "Crear ticket"),
    ("4.2", "Lista y filtros de tickets"),
    ("4.3", "Detalle, comentarios e historial"),
    ("5.1", "Agente toma ticket"),
    ("5.2", "Comentario interno"),
    ("5.3", "Resolver ticket"),
    ("6.1", "Solicitante reabre ticket"),
    ("6.2", "Cierre manual"),
    ("7.1", "Crear usuario"),
    ("7.2", "Editar usuario"),
    ("7.3", "Archivar usuario"),
    ("8.1", "Ver áreas y miembros"),
    ("8.2", "Crear área"),
    ("8.3", "Agregar miembro a área"),
    ("9.1", "Categorías — editar y toggle"),
    ("9.2", "SLAs — editar"),
    ("10.1", "Plantilla recurrente"),
    ("11.1", "Configuración del tenant"),
    ("12.1", "Reportes — Supervisor"),
    ("12.2", "Reportes — denegado a Agente"),
    ("13.1", "Notificaciones en tiempo real"),
    ("14",   "Flujo completo E2E"),
]

for i, (sec, desc) in enumerate(secciones):
    row = t_res.add_row()
    bg = "FFFFFF" if i % 2 == 0 else "F8FAFC"
    for j, txt in enumerate([sec, desc, "", "", ""]):
        set_cell_bg(row.cells[j], bg)
        row.cells[j].paragraphs[0].add_run(txt).font.size = Pt(8.5)

doc.add_paragraph()
body("Convenciones:", bold=True, size=9)
bullet("✔  Prueba exitosa — el comportamiento es el esperado")
bullet("✗   Prueba fallida — anotar el error en Observaciones")
bullet("N/A  No aplica en el entorno de pruebas actual")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  SECCIÓN 16 — NOTAS FINALES
# ════════════════════════════════════════════════════════════════════════════

heading1("16. INFORMACIÓN ADICIONAL")

heading2("Servicios activos durante las pruebas")
t_srv = doc.add_table(rows=1, cols=3)
t_srv.style = "Table Grid"
for cell, txt in zip(t_srv.rows[0].cells, ["Servicio", "URL", "Descripción"]):
    set_cell_bg(cell, "1A2C4E")
    p = cell.paragraphs[0]
    run = p.add_run(txt); run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = BLANCO
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

servicios = [
    ("Aplicativo Web",    "http://localhost:3000",  "Frontend — página principal de pruebas"),
    ("API Backend",       "http://localhost:8000/docs", "Documentación interactiva de la API (Swagger)"),
    ("Mailhog (emails)",  "http://localhost:8025",  "Visualizador de correos de notificación"),
    ("PostgreSQL",        "localhost:5432",          "Base de datos (tickets_db)"),
    ("Redis",             "localhost:6380",          "Caché y mensajería en tiempo real"),
]
for i, (svc, url, desc) in enumerate(servicios):
    row = t_srv.add_row()
    bg = "FFFFFF" if i % 2 == 0 else "F0F9FF"
    for cell, txt in zip(row.cells, [svc, url, desc]):
        set_cell_bg(cell, bg)
        cell.paragraphs[0].add_run(txt).font.size = Pt(9)

doc.add_paragraph()

heading2("Datos de prueba disponibles (seed)")
info_box(
    "El sistema fue inicializado con 20 tickets de prueba distribuidos entre diferentes estados, "
    "prioridades, áreas y categorías. Estos sirven como base para las pruebas de Dashboard y Reportes.",
    "F0FDF4", "86EFAC"
)

heading2("Verificar correos en Mailhog")
step_table([
    (1, "Abrir http://localhost:8025 en el navegador",
        "Se muestra la bandeja de Mailhog con los correos enviados por el sistema"),
    (2, "Crear un ticket o asignarlo para generar un correo",
        "El correo de notificación aparece en Mailhog con el diseño HTML del sistema"),
])

# ─ Guardar ──────────────────────────────────────────────────────────────────
output_path = "D:/Santiago/Trabajo/Smart Security/Aplicativo Tickets/Guia_Pruebas_SmartSecurity_Tickets.docx"
doc.save(output_path)
print(f"Documento generado: {output_path}")
